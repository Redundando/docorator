from docorator import Docorator
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx import Document

def main():
    # Initialize with your service account key file and email for sharing
    key_path = "service-account-key.json"
    email = "klojarve@audible.de"  # Email to share documents with
    doco = Docorator(keyfile_path="service-account-key.json", email="klojarve@audible.de", document_name="Example Document")
    doco.load()
    # Get or create a document named "Example Document"
    #document = doc_manager.get_document("Example Document")

    print(doco.url)

    document = Document()

    # Modify the document
    document.add_heading("Hello from Docorator!!!", 0)

    paragraph = document.add_paragraph("This document was created or modified using ")
    paragraph.add_run("Docorator").bold = True
    paragraph.add_run(", a Python package for Google Docs integration.")

    paragraph = document.add_paragraph("It allows you to:")
    document.add_paragraph("• Find or create Google Docs")
    document.add_paragraph("• Modify them using python-docx")
    document.add_paragraph("• Save them back to Google Drive")

    # Save the document back to Google Docs
    success = doco.save(document)

    if success:
        print("Document successfully saved!")
    else:
        print("Failed to save document.")


if __name__ == "__main__":
    main()